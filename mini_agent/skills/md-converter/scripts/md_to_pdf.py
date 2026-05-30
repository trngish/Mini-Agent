#!/usr/bin/env python3
"""
Markdown 转 PDF 转换器 - 通过 DOCX

两步转换：
1. Markdown → DOCX（使用 python-docx，中文支持良好）
2. DOCX → PDF（使用 LibreOffice 无头模式）

依赖项:
    pip install python-docx markdown
    # 还需要安装 LibreOffice 用于 PDF 转换

用法:
    from md_to_pdf import convert_md_to_pdf
    convert_md_to_pdf("input.md", "output.pdf")
"""

import os
import re
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

try:
    import markdown
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as e:
    print(f"Error: Missing dependency - {e}")
    print("Install with: uv pip install python-docx markdown")
    sys.exit(1)


def find_libreoffice():
    """查找 LibreOffice 可执行文件路径。"""
    # 常见的 Windows 路径
    windows_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice 7\program\soffice.exe",
        r"C:\Program Files\LibreOffice 8\program\soffice.exe",
    ]

    # 常见的 Linux 路径
    linux_paths = [
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "/opt/libreoffice7.0/program/soffice",
        "/opt/libreoffice7.1/program/soffice",
        "/opt/libreoffice7.2/program/soffice",
        "/opt/libreoffice7.3/program/soffice",
        "/opt/libreoffice7.4/program/soffice",
        "/opt/libreoffice7.5/program/soffice",
        "/opt/libreoffice7.6/program/soffice",
    ]

    # 检查 Windows 路径
    if sys.platform.startswith("win"):
        for path in windows_paths:
            if os.path.exists(path):
                return path
        # 尝试在 PATH 中查找
        try:
            result = subprocess.run(
                ["where", "soffice"], capture_output=True, text=True, encoding="utf-8", errors="replace"
            )
            if result.returncode == 0:
                return result.stdout.strip().split("\n")[0]
        except:
            pass

    # 检查 Linux 路径
    for path in linux_paths:
        if os.path.exists(path):
            return path

    # 在 Linux/Mac 上尝试 which 命令
    for cmd in ["soffice", "libreoffice"]:
        try:
            result = subprocess.run(["which", cmd], capture_output=True, text=True, encoding="utf-8", errors="replace")
            if result.returncode == 0:
                return result.stdout.strip()
        except:
            pass

    return None


def docx_to_pdf(docx_path, pdf_path, libreoffice_path=None):
    """使用 LibreOffice 无头模式将 DOCX 转换为 PDF。

    参数:
        docx_path: 输入 .docx 文件路径
        pdf_path: 输出 .pdf 文件路径
        libreoffice_path: 可选的 LibreOffice 可执行文件路径

    返回:
        成功返回 True，否则返回 False
    """
    if libreoffice_path is None:
        libreoffice_path = find_libreoffice()

    if libreoffice_path is None:
        raise RuntimeError(
            "LibreOffice not found. Please install LibreOffice:\n"
            "Windows: https://www.libreoffice.org/download/download/\n"
            "Linux: sudo apt install libreoffice\n"
            "Mac: brew install --cask libreoffice"
        )

    docx_path = Path(docx_path).resolve()
    pdf_path = Path(pdf_path).resolve()

    # 确保输出目录存在
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # 构建命令 - LibreOffice 无头转换
    cmd = [libreoffice_path, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)]

    try:
        subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=120)

        # LibreOffice 默认使用 --outdir 将输出到与输入相同的目录
        # 但它可能也使用相同的文件名输出
        expected_pdf = pdf_path.parent / f"{docx_path.stem}.pdf"

        if expected_pdf.exists() and expected_pdf != pdf_path:
            # 重命名为所需的输出路径
            if pdf_path.exists():
                pdf_path.unlink()
            expected_pdf.rename(pdf_path)

        return pdf_path.exists()

    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice 转换超时")
    except Exception as e:
        raise RuntimeError(f"LibreOffice 转换失败: {e}")


def process_inline_formatting(text):
    """处理内联格式并返回 (类型, 文本) 元组列表。

    返回 (类型, 文本) 元组列表，其中类型为 'text'、'bold'、'italic'、'code'。
    """
    parts = []

    # 所有内联格式的组合模式
    # 顺序很重要：先处理代码，再处理粗体，然后处理斜体
    pattern = r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)"

    last_end = 0
    for match in re.finditer(pattern, text):
        # 添加匹配前的文本
        if match.start() > last_end:
            parts.append(("text", text[last_end : match.start()]))

        matched = match.group()
        if matched.startswith("`") and matched.endswith("`"):
            parts.append(("code", matched[1:-1]))
        elif matched.startswith("**") and matched.endswith("**") or matched.startswith("__") and matched.endswith("__"):
            parts.append(("bold", matched[2:-2]))
        elif matched.startswith("*") and matched.endswith("*") or matched.startswith("_") and matched.endswith("_"):
            parts.append(("italic", matched[1:-1]))

        last_end = match.end()

    # 添加剩余文本
    if last_end < len(text):
        parts.append(("text", text[last_end:]))

    return parts if parts else [("text", text)]


def set_cell_shading(cell, color):
    """设置单元格背景颜色。"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def parse_markdown_to_docx(doc, md_content):
    """解析 markdown 内容并将元素添加到 docx 文档。"""
    lines = md_content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 标题
        if line.startswith("#######"):
            # H7 - 作为正文处理
            p = doc.add_paragraph(line[7:].strip())
        elif line.startswith("######"):
            p = doc.add_heading(line[6:].strip(), level=6)
        elif line.startswith("#####"):
            p = doc.add_heading(line[5:].strip(), level=5)
        elif line.startswith("####"):
            p = doc.add_heading(line[4:].strip(), level=4)
        elif line.startswith("###"):
            p = doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith("##"):
            p = doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith("#"):
            p = doc.add_heading(line[1:].strip(), level=1)

        # 水平线
        elif line.strip() in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p.add_run("─" * 50)

        # 块引用
        elif line.startswith(">"):
            text = line[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.italic = True

        # 无序列表
        elif line.strip().startswith(("- ", "* ", "+ ")):
            while i < len(lines) and lines[i].strip().startswith(("- ", "* ", "+ ")):
                text = lines[i].strip()[2:].strip()
                doc.add_paragraph(text, style="List Bullet")
                i += 1
            continue

        # 有序列表
        elif re.match(r"^\d+\.\s", line):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i]):
                text = re.sub(r"^\d+\.\s", "", lines[i]).strip()
                doc.add_paragraph(text, style="List Number")
                i += 1
            continue

        # 表格
        elif line.startswith("|"):
            table_data = []

            # 解析表格行
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_text = lines[i].strip().strip("|")
                if row_text and not all(c in "-: " for c in row_text):
                    cells = [c.strip() for c in row_text.split("|")]
                    table_data.append(cells)
                i += 1

            if table_data and len(table_data) >= 2:
                headers = table_data[0]
                rows = table_data[2:] if len(table_data) > 2 else []

                # 创建表格
                table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                table.style = "Table Grid"

                # 表头行
                for col_idx, header_text in enumerate(headers):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = header_text
                    # 设置表头背景颜色
                    set_cell_shading(cell, "4A90D9")
                    # 设置表头白色字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.bold = True

                # 数据行（交替颜色）
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.rows[row_idx + 1].cells[col_idx]
                        cell.text = cell_text
                        # 交替行颜色
                        if row_idx % 2 == 1:
                            set_cell_shading(cell, "F8F8F8")

                doc.add_paragraph()  # 表格后留空
            continue

        # 代码块
        elif line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1

            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            # 设置灰色背景
            p.paragraph_format.shading = True
            i += 1
            continue

        # 带内联格式的普通段落
        else:
            p = doc.add_paragraph()

            # 处理内联格式
            formatted_parts = process_inline_formatting(line)

            for part_type, text in formatted_parts:
                if not text:
                    continue

                run = p.add_run(text)

                if part_type == "bold":
                    run.bold = True
                elif part_type == "italic":
                    run.italic = True
                elif part_type == "code":
                    run.font.name = "Courier New"
                    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                    run.font.size = Pt(10)
                else:
                    # 普通文本 - 为中文设置默认字体
                    run.font.name = "Microsoft YaHei"

        i += 1


def convert_md_to_pdf(input_path: str, output_path: str, title: str = None, author: str = None, page_size: str = "A4"):
    """
    将 Markdown 文件通过 DOCX 转换为 PDF。

    两步转换：
    1. Markdown → DOCX（使用 python-docx）
    2. DOCX → PDF（使用 LibreOffice）

    参数:
        input_path: 输入 .md 文件路径
        output_path: 输出 .pdf 文件路径
        title: 文档标题（默认：文件名）
        author: 文档作者（默认："MD Converter"）
        page_size: 页面大小 - A4, Letter, Legal, A3（默认：A4）
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"输入文件未找到: {input_path}")

    # 读取 markdown 内容
    md_content = input_path.read_text(encoding="utf-8")

    # 设置默认值
    if title is None:
        title = input_path.stem
    if author is None:
        author = "MD Converter"

    # 创建文档
    doc = Document()

    # 设置文档属性
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.created = datetime.now()

    # 设置中文默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    # 如果 markdown 不以 # 开头，则添加标题
    if not md_content.strip().startswith("#"):
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 解析 markdown
    parse_markdown_to_docx(doc, md_content)

    # 创建临时 DOCX 文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        tmp_docx_path = Path(tmp_docx.name)

    try:
        # 保存 DOCX 到临时文件
        doc.save(str(tmp_docx_path))

        # 将 DOCX 转换为 PDF
        result = docx_to_pdf(tmp_docx_path, output_path)

        if result:
            print(f"✅ PDF 已创建: {output_path}")
            return str(output_path)
        else:
            raise RuntimeError("PDF 转换失败")

    finally:
        # 清理临时文件
        if tmp_docx_path.exists():
            tmp_docx_path.unlink()


def convert_md_to_docx(input_path: str, output_path: str, title: str = None, author: str = None):
    """
    将 Markdown 文件转换为 DOCX。

    参数:
        input_path: 输入 .md 文件路径
        output_path: 输出 .docx 文件路径
        title: 文档标题（默认：文件名）
        author: 文档作者（默认："MD Converter"）
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"输入文件未找到: {input_path}")

    # 读取 markdown 内容
    md_content = input_path.read_text(encoding="utf-8")

    # 设置默认值
    if title is None:
        title = input_path.stem
    if author is None:
        author = "MD Converter"

    # 创建文档
    doc = Document()

    # 设置文档属性
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.created = datetime.now()

    # 设置中文默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    # 如果 markdown 不以 # 开头，则添加标题
    if not md_content.strip().startswith("#"):
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 解析 markdown
    parse_markdown_to_docx(doc, md_content)

    # 保存
    doc.save(str(output_path))

    print(f"✅ DOCX 已创建: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="通过 DOCX 将 Markdown 转换为 PDF")
    parser.add_argument("input", help="输入 markdown 文件")
    parser.add_argument("output", help="输出 PDF 文件")
    parser.add_argument("--title", help="文档标题")
    parser.add_argument("--author", help="文档作者")
    parser.add_argument("--page-size", default="A4", choices=["A4", "Letter", "Legal", "A3"])

    args = parser.parse_args()
    convert_md_to_pdf(args.input, args.output, args.title, args.author, args.page_size)

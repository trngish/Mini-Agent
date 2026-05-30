#!/usr/bin/env python3
"""
Markdown 转 DOCX 转换器

依赖项:
    pip install python-docx markdown

用法:
    from md_to_docx import convert_md_to_docx
    convert_md_to_docx("input.md", "output.docx")
"""

import re
import sys
from datetime import datetime
from pathlib import Path

try:
    import markdown
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as e:
    print(f"Error: Missing dependency - {e}")
    print("Install with: uv pip install python-docx markdown")
    sys.exit(1)


def add_inline_formatting(run, text):
    """为 run 添加内联格式（粗体、斜体、代码）。"""
    # 代码片段（优先处理以避免冲突）
    code_pattern = r"`([^`]+)`"
    parts = re.split(code_pattern, text)

    for part in parts:
        if not part:
            continue

        if re.match(code_pattern, f"`{part}`"):
            # 代码
            run = run
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            run.font.size = Pt(10)
            run.text = part
            run = run.insert_before()
        else:
            # 检查粗体/斜体
            if part.startswith("**") and part.endswith("**") or part.startswith("__") and part.endswith("__"):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith("*") and part.endswith("*") or part.startswith("_") and part.endswith("_"):
                run.text = part[1:-1]
                run.italic = True
            else:
                run.text = part


def parse_heading(paragraph, level):
    """根据级别应用标题样式。"""
    style_name = f"Heading {level}" if level <= 6 else "Heading 2"
    paragraph.style = style_name


def parse_markdown_to_docx(doc, md_content):
    """解析 markdown 内容并将元素添加到 docx 文档。"""
    lines = md_content.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 标题
        if line.startswith("#######"):
            # H7 - 作为正文处理
            p = doc.add_paragraph(line[7:].strip())
        elif line.startswith("######"):
            p = doc.add_heading(line[6:].strip(), level=6)
        elif line.startswith("#####"):
            p = doc.add_heading(line[5:].strip(), level=5)
        elif line.startswith("####"):
            p = doc.add_heading(line[4:].strip(), level=4)
        elif line.startswith("###"):
            p = doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith("##"):
            p = doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith("#"):
            p = doc.add_heading(line[1:].strip(), level=1)

        # 水平线
        elif line.strip() in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p.add_run("─" * 50)

        # 块引用
        elif line.startswith(">"):
            text = line[1:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True

        # 无序列表
        elif line.strip().startswith(("- ", "* ", "+ ")):
            while i < len(lines) and lines[i].strip().startswith(("- ", "* ", "+ ")):
                text = lines[i].strip()[2:].strip()
                doc.add_paragraph(text, style="List Bullet")
                i += 1
            continue

        # 有序列表
        elif re.match(r"^\d+\.\s", line):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i]):
                text = re.sub(r"^\d+\.\s", "", lines[i]).strip()
                doc.add_paragraph(text, style="List Number")
                i += 1
            continue

        # 表格
        elif line.startswith("|"):
            table_data = []
            col_aligns = []

            # 解析表格行
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_text = lines[i].strip().strip("|")
                if row_text and not all(c in "-: " for c in row_text):
                    cells = [c.strip() for c in row_text.split("|")]
                    table_data.append(cells)
                elif "---" in row_text or ":--" in row_text or "--:" in row_text:
                    for c in row_text.split("|"):
                        c = c.strip()
                        if c.startswith(":") and c.endswith(":"):
                            col_aligns.append("center")
                        elif c.endswith(":"):
                            col_aligns.append("right")
                        else:
                            col_aligns.append("left")
                i += 1

            if table_data and len(table_data) >= 2:
                headers = table_data[0]
                rows = table_data[1:]

                # 创建表格
                table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                table.style = "Table Grid"

                # 表头行
                for col_idx, header_text in enumerate(headers):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = header_text
                    run = cell.paragraphs[0].runs[0]
                    run.bold = True

                # 数据行
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.rows[row_idx + 1].cells[col_idx]
                        cell.text = cell_text

                doc.add_paragraph()  # Space after table
            continue

        # 代码块
        elif line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1

            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            i += 1
            continue

        # 带内联格式的普通段落
        else:
            p = doc.add_paragraph()

            # 处理内联格式
            remaining = line

            # 按顺序处理模式
            patterns = [
                (r"\*\*([^*]+)\*\*", "bold"),
                (r"__([^_]+)__", "bold"),
                (r"\*([^*]+)\*", "italic"),
                (r"_([^_]+)_", "italic"),
                (r"`([^`]+)`", "code"),
                (r"\[([^\]]+)\]\([^\)]+\)", "link"),
            ]

            # 按模式分割文本
            parts = []
            last_end = 0

            for pattern, fmt_type in patterns:
                for match in re.finditer(pattern, remaining):
                    if match.start() > last_end:
                        parts.append(("text", remaining[last_end : match.start()]))
                    parts.append((fmt_type, match.group(1)))
                    last_end = match.end()

            if last_end < len(remaining):
                parts.append(("text", remaining[last_end:]))

            if not parts:
                parts = [("text", remaining)]

            for fmt_type, text in parts:
                if not text:
                    continue

                run = p.add_run(text)

                if fmt_type == "bold":
                    run.bold = True
                elif fmt_type == "italic":
                    run.italic = True
                elif fmt_type == "code":
                    run.font.name = "Courier New"
                    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                    run.font.size = Pt(10)
                elif fmt_type == "link":
                    run.font.color.rgb = RGBColor(0x4A, 0x90, 0xD9)
                    run.underline = True

        i += 1


def convert_md_to_docx(input_path: str, output_path: str, title: str = None, author: str = None, styles: bool = True):
    """
    将 Markdown 文件转换为 DOCX。

    参数:
        input_path: 输入 .md 文件路径
        output_path: 输出 .docx 文件路径
        title: 文档标题（默认：文件名）
        author: 文档作者（默认："MD Converter"）
        styles: 应用专业样式（默认：True）
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"输入文件未找到: {input_path}")

    # 读取 markdown 内容
    md_content = input_path.read_text(encoding="utf-8")

    # 设置默认值
    if title is None:
        title = input_path.stem
    if author is None:
        author = "MD Converter"

    # 创建文档
    doc = Document()

    # 设置文档属性
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.created = datetime.now()

    # 如果 markdown 不以 # 开头，则添加标题
    if not md_content.startswith("#"):
        doc.add_heading(title, 0)

    # 解析 markdown
    parse_markdown_to_docx(doc, md_content)

    # 保存
    doc.save(str(output_path))

    print(f"✅ DOCX created: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="将 Markdown 转换为 DOCX")
    parser.add_argument("input", help="输入 markdown 文件")
    parser.add_argument("output", help="输出 DOCX 文件")
    parser.add_argument("--title", help="文档标题")
    parser.add_argument("--author", help="文档作者")
    parser.add_argument("--no-styles", action="store_true", help="不应用样式")

    args = parser.parse_args()
    convert_md_to_docx(args.input, args.output, args.title, args.author, not args.no_styles)
